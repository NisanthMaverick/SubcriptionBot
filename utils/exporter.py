import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any
from utils.formatters import build_premium_user_details

def export_subscriptions_to_docx(subscriptions: List[Dict[str, Any]], filename: str = "subscribers_report.docx") -> str:
    """
    Exports a list of subscription records into a beautifully formatted Word document (.docx).
    """
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('💎 Premium Subscribers Report 💎', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle / Info
    sub = doc.add_paragraph(f"Total Records: {len(subscriptions)}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("────────────────────────────────────────────────────────")
    
    for idx, sub_record in enumerate(subscriptions, start=1):
        heading = doc.add_heading(f"Record #{idx}: {sub_record.get('username', 'Unknown')} (ID: {sub_record.get('user_id', '')})", level=2)
        
        # Build details text
        details_text = build_premium_user_details(sub_record)
        
        p = doc.add_paragraph()
        run = p.add_run(details_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        
        doc.add_paragraph("────────────────────────────────────────────────────────")
        
    doc.save(filename)
    return filename
